"""Экспорт викторины в DOCX."""
from io import BytesIO
from typing import Literal

from docx import Document
from docx.shared import Pt

from app.db.models import Question, Quiz
from app.services.formula_export import (
    EXPORT_FONT_BODY_PT,
    layout_rich_text,
    pptx_font_size_pt,
)
from app.services.latex_renderer import latex_render_batch
from app.services.presentation_export_service import (
    _dejavu_measure_text,
    _question_body_paragraphs,
    _resolve_formula_segment,
    _safe_filename,
    load_quiz_with_questions,
)

ExportMode = Literal["teacher", "student"]

_DIFFICULTY_LABELS = {"easy": "Лёгкая", "medium": "Средняя", "hard": "Сложная"}


def _quiz_subtitle(quiz: Quiz) -> str:
    parts = []
    if quiz.subject:
        parts.append(f"Предмет: {quiz.subject}")
    if quiz.grade:
        parts.append(f"Класс: {quiz.grade}")
    if quiz.difficulty:
        label = _DIFFICULTY_LABELS.get(quiz.difficulty, quiz.difficulty)
        parts.append(f"Сложность: {label}")
    return " · ".join(parts)


def _add_rich_text_block(doc: Document, text: str, *, level: int = 0, bullet: bool = False) -> None:
    font_size_pt = pptx_font_size_pt(level) if level else EXPORT_FONT_BODY_PT
    max_width_pt = 450.0

    layout = layout_rich_text(
        text or "",
        max_width_pt,
        font_size_pt,
        measure_text=lambda token, fs=font_size_pt: _dejavu_measure_text(token, fs),
        resolve_formula=lambda latex, display, fs=font_size_pt: _resolve_formula_segment(
            latex, display, fs,
        ),
    )

    for line in layout.lines:
        style = "List Bullet" if bullet and level > 0 else None
        paragraph = doc.add_paragraph(style=style)
        for seg in line.segments:
            if seg.kind in {"text", "fallback"}:
                if seg.content:
                    paragraph.add_run(seg.content)
                continue
            if seg.png_bytes and seg.formula_layout:
                run = paragraph.add_run()
                try:
                    run.add_picture(BytesIO(seg.png_bytes), height=Pt(seg.formula_layout.height_pt))
                except Exception:
                    paragraph.add_run(seg.content)
            else:
                paragraph.add_run(seg.content)


def build_quiz_docx(
    quiz: Quiz,
    questions: list[Question],
    mode: ExportMode,
) -> tuple[bytes, str, str]:
    doc = Document()
    doc.add_heading(quiz.title or "Викторина", 0)

    subtitle = _quiz_subtitle(quiz)
    if subtitle:
        doc.add_paragraph(subtitle)

    for index, question in enumerate(questions, start=1):
        doc.add_heading(f"Вопрос {index}", level=1)
        for text, level in _question_body_paragraphs(question, mode):
            _add_rich_text_block(doc, text, level=level, bullet=False)

    buffer = BytesIO()
    doc.save(buffer)
    filename = _safe_filename(quiz.title, "docx")
    media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    return buffer.getvalue(), filename, media_type


def export_quiz_docx(
    quiz_id: str,
    mode: ExportMode = "teacher",
) -> tuple[bytes, str, str]:
    quiz, questions = load_quiz_with_questions(quiz_id)
    with latex_render_batch():
        return build_quiz_docx(quiz, questions, mode)
